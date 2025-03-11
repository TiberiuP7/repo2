import streamlit as st
import pandas as pd
from data_processor import CAENDataProcessor
from translations import TRANSLATIONS
import io
from docx import Document

def serve_sitemap():
    with open("sitemap.xml", "r") as f:
        st.write(f.read())

if st.query_params.get("sitemap") == "1":
    serve_sitemap()

# Initialize session state
if 'language' not in st.session_state:
    st.session_state.language = 'ro'
if 'conversion_results' not in st.session_state:
    st.session_state.conversion_results = []

def get_translation(key):
    return TRANSLATIONS[st.session_state.language][key]

def load_css():
    with open('styles.css') as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def set_page_config():
    st.set_page_config(
        page_title="Convertor Coduri CAEN Rev.2 la Rev.3 | Tool Oficial 2024",
        page_icon="🔄",
        layout="wide"
    )

def initialize_data():
    if 'data_processor' not in st.session_state:
        processor = CAENDataProcessor()
        processor.extract_correspondence_data('attached_assets/Corespondenta-CAEN-Rev.2-CAEN-Rev.3.pdf')
        processor.extract_rev3_details('attached_assets/CAEN-Rev.3_structura-completa.pdf')
        st.session_state.data_processor = processor

def change_language(lang):
    st.session_state.language = lang
    st.rerun()

def export_results(results, format='excel'):
    """Export results to specified format"""
    if not results:
        return None, None, None

    df = pd.DataFrame(results)
    buffer = io.BytesIO()

    if format == 'word':
        # Create a Word document
        doc = Document()
        doc.add_heading('Rezultate Conversie CAEN', 0)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Cod CAEN Rev.2'
        header_cells[1].text = 'Cod CAEN Rev.3'
        header_cells[2].text = 'Denumire Activitate'

        for result in results:
            row_cells = table.add_row().cells
            row_cells[0].text = result['rev2_code']
            row_cells[1].text = result['rev3_code']
            row_cells[2].text = result['description']

        doc.save(buffer)
        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        file_ext = 'docx'
    else:  # Excel
        df.to_excel(buffer, index=False, engine='openpyxl')
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        file_ext = 'xlsx'

    buffer.seek(0)
    return buffer, mime, file_ext

def display_export_buttons(results, prefix=''):
    """Display export buttons and handle downloads"""
    if results:
        st.markdown(f"### 📥 {get_translation('export_results')}")
        col1, col2 = st.columns(2)

        with col1:
            word_buffer, word_mime, word_ext = export_results(results, 'word')
            if word_buffer:
                st.download_button(
                    label=f"📄 Export ca Document Word",
                    data=word_buffer,
                    file_name=f'caen_conversion_{prefix}.{word_ext}',
                    mime=word_mime,
                    key=f'word_export_{prefix}'
                )

        with col2:
            excel_buffer, excel_mime, excel_ext = export_results(results, 'excel')
            if excel_buffer:
                st.download_button(
                    label=f"📊 Export ca Excel",
                    data=excel_buffer,
                    file_name=f'caen_conversion_{prefix}.{excel_ext}',
                    mime=excel_mime,
                    key=f'excel_export_{prefix}'
                )

def display_results(results):
    """Display search results in a table format"""
    if not isinstance(results, list):
        results = [results]

    table_data = []
    for result in results:
        if isinstance(result, list):
            for r in result:
                table_data.append({
                    'Cod CAEN Rev.2': r['rev2_code'],
                    'Cod CAEN Rev.3': r['rev3_code'],
                    'Denumire Activitate': r['description']
                })
        else:
            table_data.append({
                'Cod CAEN Rev.2': result['rev2_code'],
                'Cod CAEN Rev.3': result['rev3_code'],
                'Denumire Activitate': result['description']
            })

    if table_data:
        st.markdown("### 📊 Tabel Rezultate Conversie")
        df = pd.DataFrame(table_data)
        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### Coduri CAEN Rev.2")
            st.write(df[['Cod CAEN Rev.2']].to_html(index=False), unsafe_allow_html=True)

        with col2:
            st.markdown("#### Coduri CAEN Rev.3 și Denumiri")
            df['Cod CAEN Rev.3 și Denumire'] = df['Cod CAEN Rev.3'] + ' - ' + df['Denumire Activitate']
            st.write(df[['Cod CAEN Rev.3 și Denumire']].to_html(index=False), unsafe_allow_html=True)


def validate_caen_code(code):
    """Validate a single CAEN code and return appropriate error message if invalid"""
    if not code:
        return get_translation('empty_code')
    if not code.strip().isdigit():
        return get_translation('non_numeric_code')
    code_length = len(code.strip())
    if code_length < 4:
        return get_translation('code_too_short')
    if code_length > 4:
        return get_translation('code_too_long')
    return None

def main():
    set_page_config()
    load_css()
    initialize_data()

    # Navigation menu
    st.markdown("""
        <div class="nav-container">
            <div class="nav-menu">
                <a href="/" class="active" target="_self">🔄 Convertor Coduri CAEN</a>
                <a href="/coduri_caen_rev3" target="_self">📋 Coduri CAEN Rev 3</a>
                <a href="/file_converter" target="_self">📄 Convertor Fișiere</a>
            </div>
        </div>
    """, unsafe_allow_html=True)

    # Language selector buttons
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🇷🇴", key="ro_flag", help="Romanian"):
            change_language('ro')
    with col2:
        if st.button("🇬🇧", key="en_flag", help="English"):
            change_language('en')

    # Main title with animation
    st.markdown(f'<div class="title-container"><h1>🔄 {get_translation("title")}</h1></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="subtitle-container"><h3>📋 {get_translation("subtitle")}</h3></div>', unsafe_allow_html=True)

    # Tabs for conversion options
    tab1, tab2 = st.tabs([
        get_translation("individual_conversion"),
        get_translation("bulk_conversion")
    ])

    with tab1:
        col1, col2 = st.columns([3, 1])
        with col1:
            caen_code = st.text_input(
                get_translation('search_placeholder'),
                max_chars=10,  # Increased to show error for too long inputs
                key="caen_input",
                placeholder="1234"
            )
        with col2:
            search_clicked = st.button(
                f"🔎 {get_translation('search_button')}",
                key="search_button"
            )

        if search_clicked:
            if not caen_code:
                st.error("⚠️ Vă rugăm să introduceți un cod CAEN.")
            else:
                error_message = validate_caen_code(caen_code)
                if error_message:
                    st.error(f"⚠️ {error_message}")
                else:
                    results = st.session_state.data_processor.get_conversion(caen_code)
                    if results:
                        display_results(results)
                        display_export_buttons(results, 'single')
                    else:
                        st.warning(f"{get_translation('no_results')}")

    with tab2:
        bulk_codes = st.text_area(
            get_translation('bulk_input_placeholder'),
            height=150,
            placeholder="1234\n5678\n9012"
        )
        bulk_search = st.button(
            get_translation('convert_codes'),
            key="bulk_search"
        )

        if bulk_search:
            if not bulk_codes.strip():
                st.error("⚠️ Vă rugăm să introduceți cel puțin un cod CAEN.")
            else:
                codes = [code.strip() for code in bulk_codes.split('\n') if code.strip()]
                if codes:
                    st.markdown(f"### {get_translation('conversion_results')}")
                    bulk_results = []
                    has_errors = False

                    for code in codes:
                        error_message = validate_caen_code(code)
                        if error_message:
                            st.error(f"⚠️ {error_message} pentru codul: {code}")
                            has_errors = True
                            continue

                        result = st.session_state.data_processor.get_conversion(code)
                        if result:
                            bulk_results.extend(result)
                        else:
                            st.warning(f"{get_translation('no_result_prefix')} {code}")

                    if bulk_results and not has_errors:
                        display_results(bulk_results)
                        display_export_buttons(bulk_results, 'bulk')

    # Legal reminders section
    st.markdown("---")
    st.markdown(f"""
        <div class="reminder-box">
            <p>{get_translation('important_note')}</p>
            <p>{get_translation('legal_notice')}</p>
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
