import streamlit as st
import os
from PIL import Image
import img2pdf
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from translations import TRANSLATIONS
from streamlit_js_eval import streamlit_js_eval
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextBox, LTTextLine
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf_vision_helper import extract_text_with_gpt4v
import openai

openai.api_key = api_key

def convert_pdf_to_word(pdf_file):
    """Convert PDF to Word document with improved formatting preservation"""
    try:
        # Create a new Word document
        doc = Document()
        found_text = False

        # Create BytesIO object from uploaded file
        pdf_bytes = io.BytesIO(pdf_file.read())
        pdf_bytes.seek(0)

        # Validate PDF first
        try:
            import PyPDF2
            PyPDF2.PdfReader(pdf_file)
            pdf_file.seek(0)  # Reset pointer after validation
        except Exception as e:
            st.error("Invalid or corrupted PDF file. Please ensure the file is a valid PDF document.")
            return None

        # Process each page in the PDF
        for page_layout in extract_pages(pdf_file):
            for element in page_layout:
                if isinstance(element, (LTTextBox, LTTextContainer, LTTextLine)):
                    # Create a new paragraph for each text block
                    paragraph = doc.add_paragraph()

                    # Get the text content directly from the text container
                    text_content = element.get_text()

                    # Basic formatting detection
                    font_size = 12  # default size
                    try:
                        # Try to get font size from the first character if available
                        for char in element:
                            if isinstance(char, LTChar):
                                if hasattr(char, 'size'):
                                    font_size = float(char.size)
                                break
                    except Exception as e:
                        st.warning(f"Could not determine font size for text block: {str(e)}")

                    # Handle text encoding for diacritics
                    try:
                        # Ensure proper encoding of special characters
                        text_content = text_content.encode('utf-8').decode('utf-8')
                    except UnicodeError:
                        # If encoding fails, try to clean the text
                        text_content = ''.join(char for char in text_content if ord(char) < 128)
                        st.warning(f"Some special characters might not be properly converted in paragraph: {text_content[:50]}...")

                    # Only add non-empty text content
                    if text_content.strip():
                        found_text = True
                        run = paragraph.add_run(text_content)
                        run.font.size = Pt(font_size)

                        # Add a line break after each text block if needed
                        if not text_content.endswith('\n'):
                            paragraph.add_run('\n')

            # Add a page break after each page
            doc.add_page_break()

        if not found_text:
            st.warning("Standard text extraction failed. Attempting advanced extraction...")
            try:
                # Reset file pointer before reading again
                pdf_file.seek(0)

                # Convert PDF to image bytes
                from pdf2image import convert_from_bytes
                st.info("Converting PDF pages to images...")
                pages = convert_from_bytes(pdf_file.read())
                st.success(f"Successfully converted {len(pages)} pages to images")

                # Process each page with GPT-4V
                for page_num, page in enumerate(pages, 1):
                    st.info(f"Processing page {page_num}...")
                    # Convert PIL Image to bytes
                    img_byte_arr = io.BytesIO()
                    page.save(img_byte_arr, format='PNG')
                    img_byte_arr = img_byte_arr.getvalue()

                    # Extract text using GPT-4V
                    extracted_text = extract_text_with_gpt4v(img_byte_arr)
                    if extracted_text:
                        found_text = True
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(extracted_text)
                        run.font.size = Pt(12)

                    doc.add_page_break()

                if not found_text:
                    st.error("No text could be extracted from the PDF. The file may be corrupted or password protected.")
                    return None

            except Exception as e:
                st.error(f"Advanced extraction failed: {str(e)}")
                import traceback
                st.error(f"Detailed error: {traceback.format_exc()}")
                return None

        # Save the document to a bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer.getvalue()

    except Exception as e:
        st.error(f"Error during conversion: {str(e)}")
        st.error("Please ensure the PDF is not password protected and contains extractable text")
        return None

def get_translation(key):
    return TRANSLATIONS[st.session_state.language][key]

# Initialize session state
if 'language' not in st.session_state:
    st.session_state.language = 'ro'

def change_language(lang):
    st.session_state.language = lang
    streamlit_js_eval(js_expressions=f"window.location.search = '?lang={lang}'")
    st.rerun()

def convert_image_to_pdf(image_file):
    """Convert uploaded image to PDF"""
    try:
        # Read the uploaded file
        image_bytes = image_file.read()

        # Convert to PDF using img2pdf
        pdf_bytes = img2pdf.convert(image_bytes)

        return pdf_bytes
    except Exception as e:
        st.error(f"Error during conversion: {str(e)}")
        return None

def convert_word_to_pdf(docx_file):
    """Convert uploaded Word document to PDF with proper diacritic support"""
    try:
        # Create a bytes buffer for the PDF
        pdf_buffer = io.BytesIO()

        # Read the Word document as binary
        doc_bytes = io.BytesIO(docx_file.read())
        doc = Document(doc_bytes)
        
        # Create PDF with UTF-8 encoding
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        width, height = letter

        # Register the Liberation Sans font (has good Unicode support)
        font_path = os.path.join('attached_assets', 'fonts', 'LiberationSans-Regular.ttf')
        try:
            pdfmetrics.registerFont(TTFont('LiberationSans', font_path))
            c.setFont('LiberationSans', 12)
        except Exception as e:
            st.warning("Could not load custom font for diacritics. Using fallback font.")
            c.setFont('Helvetica', 12)

        # Convert each paragraph to PDF
        y = height - 40  # Start from top of page
        for para in doc.paragraphs:
            if y < 40:  # Check if we need a new page
                c.showPage()
                try:
                    c.setFont('LiberationSans', 12)
                except:
                    c.setFont('Helvetica', 12)
                y = height - 40

            # Handle text encoding for diacritics
            text = para.text
            try:
                # Ensure proper encoding of special characters
                text = text.encode('utf-8').decode('utf-8')
            except UnicodeError:
                # If encoding fails, try to clean the text
                text = ''.join(char for char in text if ord(char) < 128)
                st.warning(f"Some special characters might not be properly converted in paragraph: {text[:50]}...")

            # Draw text with proper encoding
            c.drawString(40, y, text)
            y -= 20  # Move down for next line

        c.save()

        # Get the PDF content
        pdf_bytes = pdf_buffer.getvalue()
        return pdf_bytes

    except Exception as e:
        st.error(f"Error during conversion: {str(e)}")
        return None

def main():
    st.set_page_config(
        page_title="File Converter",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items=None
    )

    # Load custom CSS
    try:
        with open('styles.css') as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning("Style file not found. Using default styles.")

    # Navigation menu
    st.markdown("""
        <div class="nav-container">
            <div class="nav-menu">
                <a href="/" target="_self">🔄 Convertor Coduri CAEN</a>
                <a href="/coduri_caen_rev3" target="_self">📋 Coduri CAEN Rev 3</a>
                <a href="/file_converter" class="active" target="_self">📄 Convertor Fișiere</a>
            </div>
        </div>
    """, unsafe_allow_html=True)

    # Language selector buttons
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🇷🇴", key="ro_flag", help="Romanian"):
            change_language('ro')
    with col2:
        if st.button("🇬🇧", key="en_flag", help="English"):
            change_language('en')

    # Page title
    st.markdown("## 📄 Convertor Fișiere")

    # Notice about font requirement for proper diacritic support
    font_path = os.path.join('attached_assets', 'fonts', 'LiberationSans-Regular.ttf')
    if not os.path.exists(font_path):
        st.warning("""
        ⚠️ Pentru suport complet al diacriticelor în documentele Word, vă rugăm să adăugați fontul LiberationSans-Regular.ttf 
        în directorul 'attached_assets/fonts/'. Fără acest font, unele caractere speciale ar putea fi afișate incorect.
        """)

    # Create tabs for different conversion types
    image_tab, word_tab, pdf_tab = st.tabs([
        "🖼️ Image to PDF",
        "📝 Word to PDF",
        "📄 PDF to Word"
    ])

    with image_tab:
        st.markdown("### 🖼️ Conversie Imagini în PDF")
        # Image file uploader
        image_file = st.file_uploader(
            "Încărcați o imagine (JPG sau PNG)",
            type=['jpg', 'jpeg', 'png'],
            key='image_uploader'
        )

        if image_file is not None:
            # Display preview of the image
            st.image(image_file, caption='Preview imagine', use_column_width=True)

            # Convert button
            if st.button("🔄 Convertește în PDF", key='convert_image'):
                pdf_bytes = convert_image_to_pdf(image_file)

                if pdf_bytes:
                    # Create download button for the PDF
                    st.download_button(
                        label="📥 Descarcă PDF",
                        data=pdf_bytes,
                        file_name=f"{os.path.splitext(image_file.name)[0]}.pdf",
                        mime="application/pdf",
                        key='download_image_pdf'
                    )

    with word_tab:
        st.markdown("### 📝 Conversie Word în PDF")
        # Word file uploader
        word_file = st.file_uploader(
            "Încărcați un document Word (DOCX)",
            type=['docx'],
            key='word_uploader'
        )

        if word_file is not None:
            # Show file info
            st.info(f"Document încărcat: {word_file.name}")

            # Convert button
            if st.button("🔄 Convertește în PDF", key='convert_word'):
                pdf_bytes = convert_word_to_pdf(word_file)

                if pdf_bytes:
                    # Create download button for the PDF
                    st.download_button(
                        label="📥 Descarcă PDF",
                        data=pdf_bytes,
                        file_name=f"{os.path.splitext(word_file.name)[0]}.pdf",
                        mime="application/pdf",
                        key='download_word_pdf'
                    )

    with pdf_tab:
        st.markdown("### 📄 Conversie PDF în Word")

        # PDF file uploader
        pdf_file = st.file_uploader(
            "Încărcați un document PDF",
            type=['pdf'],
            key='pdf_uploader'
        )

        if pdf_file is not None:
            # Show file info
            st.info(f"Document încărcat: {pdf_file.name}")

            # Add a note about conversion quality
            st.warning("""
            ⚠️ Notă: Conversia de la PDF la Word poate să nu păstreze perfect formatarea originală.
            Vă rugăm să verificați documentul rezultat.
            """)

            # Convert button
            if st.button("🔄 Convertește în Word", key='convert_pdf'):
                try:
                    with st.spinner('Procesare document în curs...'):
                        docx_bytes = convert_pdf_to_word(pdf_file)

                    if docx_bytes:
                        st.success("Conversie finalizată cu succes!")
                        # Create download button for the Word document
                        st.download_button(
                            label="📥 Descarcă Document Word",
                            data=docx_bytes,
                            file_name=f"{os.path.splitext(pdf_file.name)[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key='download_pdf_word'
                        )
                except Exception as e:
                    st.error(f"""
                    A apărut o eroare în timpul conversiei. Asigurați-vă că:
                    - PDF-ul nu este protejat prin parolă
                    - PDF-ul conține text care poate fi extras
                    - PDF-ul nu este deteriorat

                    Detalii eroare: {str(e)}
                    """)

if __name__ == "__main__":
    main()
