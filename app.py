import streamlit as st
import pandas as pd
from data_processor import CAENDataProcessor
from translations import TRANSLATIONS
import io
from docx import Document
import streamlit.components.v1 as components
import pathlib
from bs4 import BeautifulSoup 
# asigurați-vă că `beautifulsoup4` e în requirements.txt

soup = BeautifulSoup(index_path.read_text(), 'html.parser')
noscript_tag = soup.find('noscript')
if noscript_tag:
    noscript_tag.string.replace_with("Conversie CAEN3 – convertor coduri CAEN Rev.2 la Rev.3. Vă rugăm activați JavaScript pentru a folosi aplicația.")
    index_path.write_text(str(soup))


# Path către index.html-ul Streamlit instalat
index_path = pathlib.Path(st.__file__).parent / "static" / "index.html"
# Citim conținutul actual
index_html = index_path.read_text()
# Definim meta tag-urile dorite
custom_head = """
  <title>Conversie CAEN3 – Convertor coduri CAEN Rev.2 la Rev.3</title>
  <meta name="description" content="Convertor online al codurilor CAEN (Clasificarea Activităților Economice) versiunea 2 în versiunea 3."/>
  <meta name="keywords" content="cod CAEN, conversie CAEN, CAEN Rev 2, CAEN Rev 3, clasificare, economie"/>
"""
# Injectăm meta tag-urile imediat înainte de închiderea head-ului
if custom_head not in index_html:  # evităm dublarea dacă scriptul rulează de mai multe ori
    new_html = index_html.replace('</head>', custom_head + '</head>')
    index_path.write_text(new_html)

components.html("""
    <script crossorigin="anonymous">
    // Accesăm documentul părinte al iframe-ului
    const doc = parent.document;
    // Setăm titlul paginii
    doc.title = "Conversie CAEN3 – Convertor coduri";
    // Creăm și adăugăm meta description
    let metaDesc = doc.createElement('meta');
    metaDesc.name = "description";
    metaDesc.content = "Aplicație Streamlit pentru conversia codurilor CAEN (Clasificarea Activităților Economice) de la versiunea 2 la versiunea 3.";
    doc.head.appendChild(metaDesc);
    // (Opțional) Creăm și adăugăm meta keywords
    let metaKeys = doc.createElement('meta');
    metaKeys.name = "keywords";
    metaKeys.content = "CAEN, conversie, cod CAEN Rev.2, cod CAEN Rev.3, clasificare CAEN";
    doc.head.appendChild(metaKeys);
    </script>
""", height=0)

# Initialize session state
if 'language' not in st.session_state:
    st.session_state.language = 'ro'
if 'conversion_results' not in st.session_state:
    st.session_state.conversion_results = []

def get_translation(key):
    return TRANSLATIONS[st.session_state.language][key]

def load_css():
    with open('styles.css') as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def set_page_config():
    st.set_page_config(
        page_title="Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3",
        page_icon="generated-icon.png",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items={
            'Get Help': 'https://your-help-url',
            'About': """
            ### Convertor Coduri CAEN
            Instrumentul oficial pentru conversia codurilor CAEN Rev.2 la Rev.3.
            Actualizat pentru anul 2024 conform reglementărilor în vigoare.
            """
        }
    )
    
   st.title("Conversie CAEN3 – Convertor coduri CAEN Rev.2 la Rev.3")
   st.write("Aplicație web care te ajută să convertești codurile CAEN din versiunea Rev.2 în versiunea Rev.3, oferind rapid echivalențele necesare. Ideal pentru firme și contabili care trebuie să actualizeze codurile CAEN conform noii clasificări.")


    # Enhanced SEO metadata and noscript fallback
    st.markdown(
        <head>
    <meta
      name="viewport"
      content="width=device-width, initial-scale=1, shrink-to-fit=no"
    />
    <link rel="shortcut icon" href="./favicon.png" />
    <link
      rel="preload"
      href="./static/media/SourceSansPro-Regular.DZLUzqI4.woff2"
      as="font"
      type="font/woff2"
      crossorigin
    />
    <link
      rel="preload"
      href="./static/media/SourceSansPro-SemiBold.sKQIyTMz.woff2"
      as="font"
      type="font/woff2"
      crossorigin
    />
    <link
      rel="preload"
      href="./static/media/SourceSansPro-Bold.-6c9oR8J.woff2"
      as="font"
      type="font/woff2"
      crossorigin
    />

    <title>Conversie coduri CAEN</title>

    <!-- initialize window.prerenderReady to false and then set to true in React app when app is ready for indexing -->
    <script>
      window.prerenderReady = false
    </script>
    <script type="module" crossorigin src="./static/js/index.BlylrL8P.js"></script>
    <link rel="stylesheet" crossorigin href="./static/css/index.DpJG_94W.css">
  <style media=""></style><style data-emotion="st-emotion-cache-global" data-s=""></style><style data-emotion="st-emotion-cache" data-s=""></style></head>
  <body>
    <noscript>Conversie coduri CAEN Rev. 2 la Nomenclatorul CAEN Rev.3 gratuita Actualizarea codurilor CAEN Rev. 2 la Nomenclatorul CAEN Rev. 3. Convertor Coduri CAEN Instrumentul oficial pentru conversia codurilor CAEN Rev.2 la Rev.3. Actualizat pentru anul 2024 conform reglementărilor în vigoare. Actualizeaza-ti codurile CAEN ale firmei tale la nomenclatorul CAEN Rev. 3, Tool oficial gratuit 2025. Convertor oficial și gratuit pentru codurile CAEN Rev.2 la Rev.3. Instrument actualizat 2024 pentru companii și PFA. Conversie instantă și export rezultate.</noscript>
    <div id="root"><div class=""><div class="withScreencast" data-testid="stScreencast"><div class="stApp stAppEmbeddingId-ssxmo0aj8aqd st-emotion-cache-1r4qj8v ee4bbma0" data-testid="stApp" data-test-script-state="initial" data-test-connection-state="PINGING_SERVER"><header tabindex="-1" class="stAppHeader st-emotion-cache-12fmjuu e4hpqof0" data-testid="stHeader"><div class="stDecoration st-emotion-cache-1dp5vir e4hpqof1" data-testid="stDecoration" id="stDecoration"></div><div class="stAppToolbar st-emotion-cache-15ecox0 e4hpqof2" data-testid="stToolbar"></div></header><div class="stAppViewContainer appview-container st-emotion-cache-1yiq2ps eht7o1d0" data-testid="stAppViewContainer" data-layout="narrow"><section tabindex="0" class="stMain st-emotion-cache-bm2z3a eht7o1d1" data-testid="stMain"><div class="stMainBlockContainer block-container st-emotion-cache-mtjnbi eht7o1d4" data-testid="stMainBlockContainer"><div data-testid="stVerticalBlockBorderWrapper" data-test-scroll-behavior="normal" class="st-emotion-cache-0 eu6p4el5"><div class="st-emotion-cache-b95f0i eu6p4el4"><div class="stVerticalBlock st-emotion-cache-o3een2 eu6p4el3" data-testid="stVerticalBlock" width="380"><div class="stElementContainer element-container st-emotion-cache-kj6hex eu6p4el1" data-testid="stElementContainer" data-stale="false" width="auto"><div class="stAppSkeleton st-emotion-cache-19s3t1h eu9v5nf0" data-testid="stAppSkeleton"><div class="st-emotion-cache-8hpu4s eu9v5nf1"></div><div class="st-emotion-cache-8g8ihq eu9v5nf2"><div width="98%" class="st-emotion-cache-gyzp1i eu9v5nf3"></div><div width="100%" class="st-emotion-cache-lfnwze eu9v5nf3"></div><div width="96%" class="st-emotion-cache-2xxb0l eu9v5nf3"></div><div width="65%" class="st-emotion-cache-po6pmj eu9v5nf3"></div></div><div width="75%" height="9rem" class="st-emotion-cache-faxsvr eu9v5nf4"></div></div></div></div></div></div></div><div data-testid="stAppIframeResizerAnchor" data-iframe-height="true" class="st-emotion-cache-1dumvfu eht7o1d9"></div></section></div></div></div><div data-testid="portal" id="portal" class="st-emotion-cache-1q6lfs0 e19n7mk11"></div></div><div class=""></div></div>
  

</body></html>

        """
        <head>
            <title>Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3</title>
            <meta name="description" content="Convertor oficial și gratuit pentru codurile CAEN Rev.2 la Rev.3. Instrument actualizat 2024 pentru companii și PFA. Conversie instantă și export rezultate.">
            <meta name="keywords" content="CAEN, coduri CAEN, Rev.2, Rev.3, conversie CAEN, România, 2024, actualizare CAEN, firme, PFA">
            <meta name="author" content="Convertor CAEN">
            <meta property="og:title" content="Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3">
            <meta property="og:description" content="Instrument oficial pentru conversia codurilor CAEN. Actualizat 2024, gratuit pentru companii și PFA.">
            <meta property="og:image" content="generated-icon.png">
            <meta property="og:type" content="website">
            <noscript>
                <h1>Convertor Coduri CAEN 2024</h1>
                <p>Instrument oficial pentru conversia codurilor CAEN Rev.2 la Rev.3. Vă rugăm să activați JavaScript pentru a utiliza aplicația.</p>
            </noscript>
        </head>
    """, unsafe_allow_html=True)

    st.html("<head><title>Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3</title></head>")
    st.html("<head><meta name="description" content="Convertor oficial și gratuit pentru codurile CAEN Rev.2 la Rev.3. Instrument actualizat 2024 pentru companii și PFA. Conversie instantă și export rezultate."></head>")
    st.html("<head><meta name="keywords" content="CAEN, coduri CAEN, Rev.2, Rev.3, conversie CAEN, România, 2024, actualizare CAEN, firme, PFA"></head>")
    st.html("<head><meta name="author" content="Convertor CAEN"></head>")
    st.html("<head><meta property="og:title" content="Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3"></head>")
    st.html("<head><meta property="og:description" content="Instrument oficial pentru conversia codurilor CAEN. Actualizat 2024, gratuit pentru companii și PFA."></head>")
    st.html("<head><meta property="og:image" content="generated-icon.png"></head>")
    st.html("<head><meta property="og:type" content="website"></head>")
            
@st.cache_resource
def get_data_processor():
    processor = CAENDataProcessor()
    processor.extract_correspondence_data('attached_assets/Corespondenta-CAEN-Rev.2-CAEN-Rev.3.pdf')
    processor.extract_rev3_details('attached_assets/CAEN-Rev.3_structura-completa.pdf')
    return processor

def initialize_data():
    if 'data_processor' not in st.session_state:
        st.session_state.data_processor = get_data_processor()

def change_language(lang):
    st.session_state.language = lang
    st.rerun()

def export_results(results, format='excel'):
    """Export results to specified format"""
    if not results:
        return None, None, None

    df = pd.DataFrame(results)
    buffer = io.BytesIO()

    if format == 'word':
        # Create a Word document
        doc = Document()
        doc.add_heading('Rezultate Conversie CAEN', 0)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Cod CAEN Rev.2'
        header_cells[1].text = 'Cod CAEN Rev.3'
        header_cells[2].text = 'Denumire Activitate'

        for result in results:
            row_cells = table.add_row().cells
            row_cells[0].text = result['rev2_code']
            row_cells[1].text = result['rev3_code']
            row_cells[2].text = result['description']

        doc.save(buffer)
        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        file_ext = 'docx'
    else:  # Excel
        df.to_excel(buffer, index=False, engine='openpyxl')
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        file_ext = 'xlsx'

    buffer.seek(0)
    return buffer, mime, file_ext

def display_export_buttons(results, prefix=''):
    """Display export buttons and handle downloads"""
    if results:
        st.markdown(f"### 📥 {get_translation('export_results')}")
        col1, col2 = st.columns(2)

        with col1:
            word_buffer, word_mime, word_ext = export_results(results, 'word')
            if word_buffer:
                st.download_button(
                    label=f"📄 Export ca Document Word",
                    data=word_buffer,
                    file_name=f'caen_conversion_{prefix}.{word_ext}',
                    mime=word_mime,
                    key=f'word_export_{prefix}'
                )

        with col2:
            excel_buffer, excel_mime, excel_ext = export_results(results, 'excel')
            if excel_buffer:
                st.download_button(
                    label=f"📊 Export ca Excel",
                    data=excel_buffer,
                    file_name=f'caen_conversion_{prefix}.{excel_ext}',
                    mime=excel_mime,
                    key=f'excel_export_{prefix}'
                )

def display_results(results):
    """Display search results in a table format"""
    if not isinstance(results, list):
        results = [results]

    table_data = []
    for result in results:
        if isinstance(result, list):
            for r in result:
                table_data.append({
                    'Cod CAEN Rev.2': r['rev2_code'],
                    'Cod CAEN Rev.3': r['rev3_code'],
                    'Denumire Activitate': r['description']
                })
        else:
            table_data.append({
                'Cod CAEN Rev.2': result['rev2_code'],
                'Cod CAEN Rev.3': result['rev3_code'],
                'Denumire Activitate': result['description']
            })

    if table_data:
        st.markdown("### 📊 Tabel Rezultate Conversie")
        df = pd.DataFrame(table_data)
        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### Coduri CAEN Rev.2")
            st.write(df[['Cod CAEN Rev.2']].to_html(index=False), unsafe_allow_html=True)

        with col2:
            st.markdown("#### Coduri CAEN Rev.3 și Denumiri")
            df['Cod CAEN Rev.3 și Denumire'] = df['Cod CAEN Rev.3'] + ' - ' + df['Denumire Activitate']
            st.write(df[['Cod CAEN Rev.3 și Denumire']].to_html(index=False), unsafe_allow_html=True)


def validate_caen_code(code):
    """Validate a single CAEN code and return appropriate error message if invalid"""
    if not code:
        return get_translation('empty_code')
    if not code.strip().isdigit():
        return get_translation('non_numeric_code')
    code_length = len(code.strip())
    if code_length < 4:
        return get_translation('code_too_short')
    if code_length > 4:
        return get_translation('code_too_long')
    return None

def main():
    set_page_config()
    load_css()
    initialize_data()

     st.markdown("""
        <head>
            <title>Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3</title>
            <meta name="description" content="Convertor oficial și gratuit pentru codurile CAEN Rev.2 la Rev.3. Instrument actualizat 2024 pentru companii și PFA. Conversie instantă și export rezultate.">
            <meta name="keywords" content="CAEN, coduri CAEN, Rev.2, Rev.3, conversie CAEN, România, 2024, actualizare CAEN, firme, PFA">
            <meta name="author" content="Convertor CAEN">
            <meta property="og:title" content="Convertor Coduri CAEN 2024 | Conversie Gratuită Rev.2 la Rev.3">
            <meta property="og:description" content="Instrument oficial pentru conversia codurilor CAEN. Actualizat 2024, gratuit pentru companii și PFA.">
            <meta property="og:image" content="generated-icon.png">
            <meta property="og:type" content="website">
            <noscript>
                <h1>Convertor Coduri CAEN 2024</h1>
                <p>Instrument oficial pentru conversia codurilor CAEN Rev.2 la Rev.3. Vă rugăm să activați JavaScript pentru a utiliza aplicația.</p>
            </noscript>
        </head>
    """, unsafe_allow_html=True)
    
    # Navigation menu
    st.markdown("""
        <div class="nav-container">
            <div class="nav-menu">
                <a href="/" class="active" target="_self">🔄 Convertor Coduri CAEN</a>
                <a href="/coduri_caen_rev3" target="_self">📋 Coduri CAEN Rev 3</a>
                <a href="/file_converter" target="_self">📄 Convertor Fișiere</a>
            </div>
        </div>
    """, unsafe_allow_html=True)

    # Language selector buttons
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🇷🇴", key="ro_flag", help="Romanian"):
            change_language('ro')
    with col2:
        if st.button("🇬🇧", key="en_flag", help="English"):
            change_language('en')

    # Main title with animation
    st.markdown(f'<div class="title-container"><h1>🔄 {get_translation("title")}</h1></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="subtitle-container"><h3>📋 {get_translation("subtitle")}</h3></div>', unsafe_allow_html=True)

    # Tabs for conversion options
    tab1, tab2 = st.tabs([
        get_translation("individual_conversion"),
        get_translation("bulk_conversion")
    ])

    with tab1:
        col1, col2 = st.columns([3, 1])
        with col1:
            caen_code = st.text_input(
                get_translation('search_placeholder'),
                max_chars=10,  # Increased to show error for too long inputs
                key="caen_input",
                placeholder="1234"
            )
        with col2:
            search_clicked = st.button(
                f"🔎 {get_translation('search_button')}",
                key="search_button"
            )

        if search_clicked:
            if not caen_code:
                st.error("⚠️ Vă rugăm să introduceți un cod CAEN.")
            else:
                error_message = validate_caen_code(caen_code)
                if error_message:
                    st.error(f"⚠️ {error_message}")
                else:
                    results = st.session_state.data_processor.get_conversion(caen_code)
                    if results:
                        display_results(results)
                        display_export_buttons(results, 'single')
                    else:
                        st.warning(f"{get_translation('no_results')}")

    with tab2:
        bulk_codes = st.text_area(
            get_translation('bulk_input_placeholder'),
            height=150,
            placeholder="1234\n5678\n9012"
        )
        bulk_search = st.button(
            get_translation('convert_codes'),
            key="bulk_search"
        )

        if bulk_search:
            if not bulk_codes.strip():
                st.error("⚠️ Vă rugăm să introduceți cel puțin un cod CAEN.")
            else:
                codes = [code.strip() for code in bulk_codes.split('\n') if code.strip()]
                if codes:
                    st.markdown(f"### {get_translation('conversion_results')}")
                    bulk_results = []
                    has_errors = False

                    for code in codes:
                        error_message = validate_caen_code(code)
                        if error_message:
                            st.error(f"⚠️ {error_message} pentru codul: {code}")
                            has_errors = True
                            continue

                        result = st.session_state.data_processor.get_conversion(code)
                        if result:
                            bulk_results.extend(result)
                        else:
                            st.warning(f"{get_translation('no_result_prefix')} {code}")

                    if bulk_results and not has_errors:
                        display_results(bulk_results)
                        display_export_buttons(bulk_results, 'bulk')

    # Legal reminders section
    st.markdown("---")
    st.markdown(f"""
        <div class="reminder-box">
            <p>{get_translation('important_note')}</p>
            <p>{get_translation('legal_notice')}</p>
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
